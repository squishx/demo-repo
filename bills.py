from docx import Document


# Création du document Word
doc = Document()
doc.add_heading('SEEG', 0)

# Ajout de la date et heure
doc.add_paragraph('17/05/25\t\t17:49')

# Infos principales
doc.add_paragraph('No de Reçu : 101047/2528957')
doc.add_paragraph('No du Compteur : 01327564223')
doc.add_paragraph('Prépaiement Electricité')
doc.add_paragraph('MR MOUELE OGBUH QUISH LIONEL SAMUEL')
doc.add_paragraph('7144310425')

doc.add_paragraph('Crédit Service')
doc.add_paragraph('Crédit d\'électricité')

# Numéro Jeton
doc.add_paragraph('Numéro Jeton :')
doc.add_paragraph('1910 6360 5967')
doc.add_paragraph('7635 9514')

# Infos supplémentaires
doc.add_paragraph('SGC : 600001')
doc.add_paragraph('TI : 03')
doc.add_paragraph('KRN : 2')

# Montant
doc.add_paragraph('Montant encaissé : 2000 XAF')
doc.add_paragraph('Total Unités (kwh) : 15,8')
doc.add_paragraph('Valeur Totale : 2000 XAF')

# Détails
doc.add_paragraph('BT 3 kW : 1742')
doc.add_paragraph('@ 111.18/kwh')
doc.add_paragraph('C.S.E Réduite : 93')
doc.add_paragraph('C.O.M Réduite @ 3.5% : 64')
doc.add_paragraph('C.S.S Réduite @ 0.5% : 9')
doc.add_paragraph('T.V.A Réduite @ 5.0% : 92')

# Footer
doc.add_paragraph('Opérateur : SEEG KPOWERS')
doc.add_paragraph('Revendeur : ALLES')

# Sauvegarder le fichier
file_path = '/mnt/data/credit_electricite_modifiable.docx'
doc.save(file_path)

file_path = '/mnt/data/credit_electricite_modifiable.docx'